""" This program will create a word document with 3 random taco recipes and a randomly selected taco image. """
# Importing necessary libraries for program
import requests  # Importing requests for API calls
import docx  # Allows for creation of Microsoft Word documents
# Importing Pillow image manipulation library and modules for writing on the images
from PIL import Image, ImageFont, ImageDraw
from docx.shared import Inches  # Importing the Inches module to use when sizing the image in the word doc
import shutil  # Lets me save an image from a url

"""
In the following block, I'll be resizing, and adding a title to the taco image. I will then add the image and 
credits to the first page of the document
"""
# Try to open the taco image. If it's not in your working directory, an error message will print.
try:
    """Making an API call to unsplash for a random image with the query parameter "taco" and my unique client_id. 
    Data is converted to json formatting."""
    random_taco_image = requests.get('https://api.unsplash.com/photos/random?query="taco";client_id=3a31301fc8edce2497d0e0dd001ec1bfbed7a207b9207dd2ccc5c87baade2b12').json()
    # Parsing the json data to retrieve the taco image url
    image_url = random_taco_image['links']['download']
    """
    The next 5 lines of code have been adapted from: 
    'https://www.dev2qa.com/how-to-download-image-file-from-url-use-python-requests-or-wget-module/'
    and their purpose is to download an image from a given url.
    """
    # Open the url image, set stream to True, this will return the stream content.
    resp = requests.get(image_url, stream=True)
    # Open a local file with wb ( write binary ) permission.
    local_file = open('taco_thumbnail.jpg', 'wb')
    # Set decode_content value to True, otherwise the downloaded image file's size will be zero.
    resp.raw.decode_content = True
    # Copy the response stream raw data to local image file.
    shutil.copyfileobj(resp.raw, local_file)
    # Remove the image url response object.
    del resp
    # If the file is misplaced and is not in the working directory, an error message will print.
    try:
        image = Image.open('taco_thumbnail.jpg')
        # Trying to resize the image. If something goes wrong, an error message will print.
        try:
            image.thumbnail((800, 800))  # Resizing the image to a thumbnail
            # Getting the font for the image text and choosing it's size
            font = ImageFont.truetype('DejaVuSans.ttf', 45)
            img_draw = ImageDraw.Draw(image)  # Creating new object to draw on image
            # Defining the look and position of text on the image
            img_draw.text([10, 50], 'Random Taco Cookbook', fill='white', font=font)
            image.save('taco_thumbnail.jpg')  # Saving the image with a new file name
        # One time when I ran this program I was given this error because the image was truncated
        except OSError:
            # Because the API is returning random images, there are resizing problems occasionally.
            print('There was a problem resizing your image. Try running the program again.')

    # Try to begin creating a word document. If the document is currently open. An error message will print
        try:
            # This portion of code is focused on creating the first page of the Word document
            document = docx.Document()  # Creating a Word document
            document.add_paragraph('Random Taco Cookbook', 'Title')  # Adding a title to first page with the 'Title' style
            # Adding the taco image with a width of 5.5 inches
            document.add_picture('taco_thumbnail.jpg', width=Inches(5.5))
            document.add_paragraph('Credits', 'Heading 1')  # Adding a paragraph with the word 'credits'
            # This paragraph gives credit to the person who took the photo
            document.add_paragraph('•  Taco image from: https://api.unsplash.com')
            # This paragraph gives credit to the taco recipes API
            document.add_paragraph('•  Taco recipes from: https://taco-1150.herokuapp.com/random/?full_taco=true')
            # This paragraph states my name as the coder for this program
            document.add_paragraph('•  Code by Andrea Pratt')
            document.add_page_break()
            document.save('Random_Taco_Cookbook.docx')  # Saving the document before trying to call the API
        # Try to call the API and add the data to word document. If it fails there will be an error message
            try:
                # This loop will be responsible for creating 3 separate and random taco recipes
                for number in range(3):
                    # Calling the taco recipe API (once for each recipe)
                    data = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()
                    """For each of these recipes, I'm retrieving the 'recipe' key and the 'name' key
                    from the API call and saving each one in a variable named accordingly"""
                    base_layer_recipe = data['base_layer']['recipe']
                    seasoning_recipe = data['seasoning']['recipe']
                    mixin_recipe = data['mixin']['recipe']
                    condiment_recipe = data['condiment']['recipe']
                    shell_recipe = data['shell']['recipe']
                    base_layer_name = data['base_layer']['name']
                    seasoning_name = data['seasoning']['name']
                    mixin_name = data['mixin']['name']
                    condiment_name = data['condiment']['name']
                    shell_name = data['shell']['name']
                    # Here I'm adding a header for each of the taco recipes that lists all 5 taco components
                    document.add_paragraph(f'{base_layer_name} with {seasoning_name}, {mixin_name} and {condiment_name} '
                                           f'in {shell_name}', 'Heading 1')
                    # Adding taco base layer name with the heading 1 style
                    document.add_paragraph(f'{base_layer_name}', 'Heading 4')
                    # Adding a paragraph that is the base layer recipe
                    document.add_paragraph(f'{base_layer_recipe}')
                    # Adding a heading paragraph for the seasoning recipe name
                    document.add_paragraph(f'{seasoning_name}', 'Heading 4')
                    # Adding a paragraph for the seasoning recipe
                    document.add_paragraph(f'{seasoning_recipe}')
                    # Adding a heading paragraph with the mixin recipe name
                    document.add_paragraph(f'{mixin_name}', 'Heading 4')
                    # Adding a paragraph with the mixin recipe
                    document.add_paragraph(f'{mixin_recipe}')
                    # Adding a header paragraph for the condiment name
                    document.add_paragraph(f'{condiment_name}', 'Heading 4')
                    # Adding a paragraph for the condiment recipe
                    document.add_paragraph(f'{condiment_recipe}')
                    # Adding a header paragraph for the shell name
                    document.add_paragraph(f'{shell_name}', 'Heading 4')
                    # Adding a paragraph for the shell recipe
                    document.add_paragraph(f'{shell_recipe}')
                    # Adding a page break between each of the 3 recipes, but not at the end of the document
                    if number != 2:
                        document.add_page_break()
                document.save('Random_Taco_Cookbook.docx')
                print('Your random taco cookbook has been created successfully!')
                print('Open the Random_Taco_Cookbook.docx file to view!')
        # If the internet connection fails, this error message will print
            except requests.exceptions.ConnectionError:
                print('There seems to be a problem with your internet connection. Cannot get recipe data.')

    # If the word document is already opening when running this code, this error message will print
        except PermissionError:
            print('Please close the existing word document and try again.')
# if the taco image file isn't in your working directory, this error message will print
    except FileNotFoundError:
        print('The image file you are trying to open is not in your working directory.')
# If the taco image API fails, it will print an error message telling the user to check their internet connection
except requests.exceptions.ConnectionError:
    print('There seems to be a problem with your internet connection. Cannot get recipe data.')